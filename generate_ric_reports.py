#!/usr/bin/env python3
"""
Generate RIC PM Reports - May 2026
Creates a comprehensive Word document with 7 PM reports.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
DARK_BLUE = RGBColor(0, 70, 127)
WHITE     = RGBColor(255, 255, 255)
LIGHT_GRAY = RGBColor(242, 242, 242)
RED_COLOR  = RGBColor(192, 0, 0)
BODY_FONT  = "Calibri"
BODY_SIZE  = Pt(11)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_background(cell, hex_color: str):
    """Set cell shading/background fill."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    """Add a page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type='page')
    # Remove paragraph spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def rgb_to_hex(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


DARK_BLUE_HEX  = rgb_to_hex(0, 70, 127)
LIGHT_GRAY_HEX = rgb_to_hex(242, 242, 242)
WHITE_HEX      = rgb_to_hex(255, 255, 255)
RED_HEX        = rgb_to_hex(255, 235, 235)  # light red for blocked rows


def style_paragraph(para, size=Pt(11), bold=False, color=None, align=None, space_before=None, space_after=None):
    """Apply styling to all runs in a paragraph, or to the para itself."""
    fmt = para.paragraph_format
    if align:
        para.alignment = align
    if space_before is not None:
        fmt.space_before = space_before
    if space_after is not None:
        fmt.space_after = space_after
    for run in para.runs:
        run.font.name = BODY_FONT
        run.font.size = size
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """Add a dark-blue section heading."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    return para


def add_subheading(doc, text):
    """Smaller dark-blue subheading."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(3)
    return para


def add_body(doc, text, bold=False, italic=False, color=None):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(11)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    return para


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point paragraph."""
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.bold = True
        r2 = para.add_run(text)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(11)
    else:
        r = para.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    return para


def make_table(doc, headers, rows, col_widths=None, blocked_col=None, blocked_values=None):
    """
    Create a styled table.
    headers: list of str
    rows: list of list of str
    col_widths: list of Inches (optional)
    blocked_col: index of column to check for 'blocked' highlight
    blocked_values: set of values that trigger red highlight
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_background(cell, DARK_BLUE_HEX)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.name  = BODY_FONT
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        is_gray = (row_idx % 2 == 1)
        # Check if blocked
        is_blocked = False
        if blocked_col is not None and blocked_values:
            val = row_data[blocked_col] if blocked_col < len(row_data) else ""
            if any(bv.lower() in val.lower() for bv in blocked_values):
                is_blocked = True

        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            if is_blocked:
                set_cell_background(cell, RED_HEX)
            elif is_gray:
                set_cell_background(cell, LIGHT_GRAY_HEX)
            else:
                set_cell_background(cell, WHITE_HEX)

            p = cell.paragraphs[0]
            p.clear()
            # Color blocked text red
            if is_blocked and "BLOCKED" in cell_text.upper():
                run = p.add_run(cell_text)
                run.font.color.rgb = RED_COLOR
                run.font.bold = True
            else:
                run = p.add_run(cell_text)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_report_header(doc, text):
    """Add a dark-blue banner paragraph for report page header."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = WHITE
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(12)
    # Shade paragraph background
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), DARK_BLUE_HEX)
    pPr.append(shd)
    # Add padding via paragraph border / spacing — use indent
    para.paragraph_format.left_indent  = Inches(0.1)
    return para


def add_footer(doc):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RIC PM Reports  |  eSimplicity  |  May 13, 2026  |  Confidential")
        run.font.name  = BODY_FONT
        run.font.size  = Pt(9)
        run.font.color.rgb = DARK_BLUE
        run.font.italic = True


# ─────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────

def build_cover_page(doc):
    # Spacer
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RIC Project Management Reports")
    run.font.name  = BODY_FONT
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
    title_para.paragraph_format.space_after = Pt(10)

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run("Rapid Innovation Center  |  eSimplicity  |  May 13, 2026")
    run2.font.name  = BODY_FONT
    run2.font.size  = Pt(14)
    run2.font.color.rgb = DARK_BLUE
    run2.font.italic = True
    sub_para.paragraph_format.space_after = Pt(30)

    # Divider line
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), DARK_BLUE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    div.paragraph_format.space_after = Pt(20)

    # Report list heading
    list_head = doc.add_paragraph()
    list_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = list_head.add_run("Contents")
    rh.font.name  = BODY_FONT
    rh.font.size  = Pt(13)
    rh.font.bold  = True
    rh.font.color.rgb = DARK_BLUE
    list_head.paragraph_format.space_after = Pt(8)

    reports = [
        "1.  Daily Morning Snapshot",
        "2.  Board Health Report",
        "3.  Team Workload Report",
        "4.  Risk & Blockers Report",
        "5.  Monthly Executive Summary",
        "6.  Domain & Label Distribution Report",
        "7.  Initiative Progress Report",
    ]
    for r in reports:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(r)
        run.font.name  = BODY_FONT
        run.font.size  = Pt(12)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # Bottom note
    for _ in range(4):
        doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("RIC PM Reports  |  eSimplicity  |  May 13, 2026  |  Confidential")
    nr.font.name   = BODY_FONT
    nr.font.size   = Pt(9)
    nr.font.italic = True
    nr.font.color.rgb = DARK_BLUE


# ─────────────────────────────────────────────
# REPORT 1: Daily Morning Snapshot
# ─────────────────────────────────────────────

def build_report1(doc):
    add_report_header(doc, "1. Daily Morning Snapshot — May 13, 2026")

    add_heading(doc, "What Moved Yesterday (May 12)")
    add_body(doc, "New tickets created May 12, 2026:")
    new_tickets = [
        ("RIC-417", "Create Apex Package - CIO Growth", "Initiative", "In Progress", "Kyle Hogan"),
        ("RIC-431", "NHLBI - Apex Agent Assistance", "Initiative", "Discovery", "Unassigned"),
        ("RIC-436", "eSim Partnership Management", "Initiative", "Always On", "Kyle Hogan"),
        ("RIC-457", "SHINE MVP - Milestone 5", "Epic", "Backlog", "Unassigned"),
        ("RIC-458", "SHINE MVP - Milestone 6", "Epic", "Backlog", "Unassigned"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Status", "Owner"],
               [[t[0], t[1], t[2], t[3], t[4]] for t in new_tickets],
               col_widths=[Inches(0.9), Inches(2.8), Inches(0.9), Inches(1.0), Inches(1.2)])

    add_heading(doc, "Currently Blocked — Needs Attention Today")
    blocked = [
        ("RIC-261", "Create Apex Package - T-MSIS", "Initiative", "Kyle Hogan", "~29 days"),
        ("RIC-251", "Security Use Case 2 - SIA Automation", "Epic", "Matt Wartell", "~35 days"),
        ("RIC-250", "Security Use Case 1 - Scans Review", "Epic", "Matt Wartell", "~35 days"),
        ("RIC-185", "Collect FOE Success Story from OPERAS", "Story", "Aaron West", "~51 days"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Owner", "Days Blocked"],
               [list(b) for b in blocked],
               col_widths=[Inches(0.9), Inches(2.8), Inches(0.9), Inches(1.1), Inches(1.1)],
               blocked_col=0,
               blocked_values={"RIC-261", "RIC-251", "RIC-250", "RIC-185"})

    add_heading(doc, "Currently Waiting — External Dependencies")
    waiting = [
        ("RIC-274", "Implement Context Extraction Module", "Matt Wartell"),
        ("RIC-188", "Collect FOE Success Story from USGS", "Aaron West"),
        ("RIC-407", "Get overview of SIA process from client", "Matt Wartell"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Owner"],
               [list(w) for w in waiting],
               col_widths=[Inches(0.9), Inches(3.8), Inches(1.5)])

    add_heading(doc, "Ready to Ship — Sitting in Ready for QA")
    ready = [
        ("RIC-353", "Diagramming Agent — architecture and flow diagram skills", "Fareez Ahmed"),
        ("RIC-419", "Package and deliver Apex proposal agents to deployment team", "Thomas Mason"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Owner"],
               [list(r) for r in ready],
               col_widths=[Inches(0.9), Inches(3.8), Inches(1.5)])

    add_heading(doc, "Items in Review")
    reviewing = [
        ("RIC-358", "Review T-MSIS Apex Agent", "Thomas Mason"),
        ("RIC-423", "Add BDD post-deploy verification to prod", "Thomas Mason"),
        ("RIC-268", "IDR - Research POC/White Paper (Initiative)", "Thomas Mason"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Owner"],
               [list(r) for r in reviewing],
               col_widths=[Inches(0.9), Inches(3.8), Inches(1.5)])


# ─────────────────────────────────────────────
# REPORT 2: Board Health Report
# ─────────────────────────────────────────────

def build_report2(doc):
    add_report_header(doc, "2. Board Health Report — May 13, 2026")
    add_body(doc, "Flow health across all three RIC boards. Focus areas: stale work, unassigned items, and bottlenecks.")

    add_heading(doc, "Status Distribution (Verified)")

    # Three side-by-side tables using a container table
    container = doc.add_table(rows=1, cols=3)
    container.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove container table borders
    tbl = container._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    def mini_status_table(cell, title, total, data):
        """Build a mini status count table inside a cell."""
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(f"{title} ({total} total)")
        run.font.name  = BODY_FONT
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_after = Pt(4)

        tbl = cell.add_table(rows=1 + len(data), cols=2)
        tbl.style = 'Table Grid'

        # Header
        hcells = tbl.rows[0].cells
        for i, h in enumerate(["Status", "Count"]):
            set_cell_background(hcells[i], DARK_BLUE_HEX)
            p2 = hcells[i].paragraphs[0]
            p2.clear()
            r = p2.add_run(h)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = WHITE
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data
        for ri, (status, count) in enumerate(data):
            rcells = tbl.rows[ri+1].cells
            is_gray = (ri % 2 == 1)
            for ci, val in enumerate([status, str(count)]):
                if is_gray:
                    set_cell_background(rcells[ci], LIGHT_GRAY_HEX)
                else:
                    set_cell_background(rcells[ci], WHITE_HEX)
                p3 = rcells[ci].paragraphs[0]
                p3.clear()
                r = p3.add_run(val)
                r.font.name = BODY_FONT
                r.font.size = Pt(9)
                if ci == 1:
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if status.lower() == "blocked":
                    r.font.color.rgb = RED_COLOR
                    r.font.bold = True

    initiatives_data = [
        ("In Progress", 6), ("Always On", 4), ("To Do", 4), ("Discovery", 3),
        ("Ready for Work", 2), ("Reviewing", 1), ("Blocked", 1),
    ]
    epics_data = [
        ("In Progress", 20), ("To Do", 15), ("Blocked", 2),
    ]
    team_data = [
        ("In Progress", 20), ("To Do", 6), ("Waiting", 3),
        ("Ready for QA", 2), ("Reviewing", 1), ("Blocked", 1),
    ]

    cells = container.rows[0].cells
    mini_status_table(cells[0], "INITIATIVES", 21, initiatives_data)
    mini_status_table(cells[1], "EPICS", 37, epics_data)
    mini_status_table(cells[2], "TEAM BOARD", 33, team_data)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading(doc, "Stale Work — In Progress > 30 Days")
    add_subheading(doc, "Over 60 Days")
    stale_60 = [
        ("RIC-9",   "CQP Training Material Automation",         "Initiative", "Unassigned",       "96 days", "May 8",  ""),
        ("RIC-10",  "QDAS - AI Acceleration",                   "Initiative", "Unassigned",       "96 days", "May 8",  ""),
        ("RIC-142", "Agent Orchestration Platform Work",         "Epic",       "Aaron West",       "63 days", "Apr 7",  ""),
        ("RIC-159", "NTIA Reports Site Deployment",              "Epic",       "Thomas Mason",     "56 days", "Apr 27", ""),
        ("RIC-184", "Collect FOE Success Story QDAS",            "Story",      "Aaron West",       "51 days", "May 6",  ""),
        ("RIC-185", "Collect FOE Success Story OPERAS",          "Story",      "Aaron West",       "51 days", "May 6",  "BLOCKED"),
        ("RIC-188", "Collect FOE Success Story USGS",            "Story",      "Aaron West",       "51 days", "May 6",  "WAITING"),
        ("RIC-189", "Define Website User Types",                 "Story",      "Aaron West",       "51 days", "May 6",  ""),
        ("RIC-190", "Build FOE Maturity Quiz",                   "Story",      "Aaron West",       "51 days", "May 11", ""),
        ("RIC-199", "Finish FOE Presentation Deck",              "Story",      "Aaron West",       "48 days", "May 6",  ""),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Owner", "Days In Progress", "Last Updated", "Flag"],
               [[s[0], s[1], s[2], s[3], s[4], s[5], s[6]] for s in stale_60],
               col_widths=[Inches(0.75), Inches(2.4), Inches(0.7), Inches(1.0), Inches(0.9), Inches(0.85), Inches(0.7)],
               blocked_col=6,
               blocked_values={"BLOCKED"})

    add_subheading(doc, "Over 30 Days")
    stale_30 = [
        ("RIC-225", "Katalyst",                             "Epic",   "Thomas Mason",   "37 days", "Apr 17", ""),
        ("RIC-223", "Prove Value of AI Agents NHLBI",       "Story",  "Chris Pusinelli", "37 days", "May 12", ""),
        ("RIC-274", "Implement Context Extraction Module",  "Story",  "Matt Wartell",    "23 days", "May 11", "WAITING"),
        ("RIC-198", "Get ELT Sign-Off for FOE",             "Epic",   "Aaron West",      "48 days", "Apr 27", ""),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Owner", "Days In Progress", "Last Updated", "Flag"],
               [[s[0], s[1], s[2], s[3], s[4], s[5], s[6]] for s in stale_30],
               col_widths=[Inches(0.75), Inches(2.4), Inches(0.7), Inches(1.0), Inches(0.9), Inches(0.85), Inches(0.7)])

    add_heading(doc, "Unassigned Items by Board")
    unassigned = [
        ("Initiatives", "7", "RIC-9, RIC-10, RIC-88, RIC-133, RIC-145, RIC-365, RIC-431"),
        ("Epics",        "8", "RIC-175, RIC-176, RIC-177, RIC-179, RIC-181, RIC-182, RIC-318, RIC-319"),
        ("Team Board",   "1", "RIC-427 (AI Gateway Infrastructure Foundation)"),
    ]
    make_table(doc,
               ["Board", "Unassigned Count", "Notable Items"],
               [list(u) for u in unassigned],
               col_widths=[Inches(1.2), Inches(1.3), Inches(4.3)])

    add_heading(doc, "Bottleneck Signal — Ready for QA Not Moving")
    bottleneck = [
        ("RIC-353", "Diagramming Agent — architecture and flow diagram skills", "Fareez Ahmed",  "Sitting since ~May 8"),
        ("RIC-419", "Package and deliver Apex proposal agents to deployment team", "Thomas Mason", "Created May 12"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Owner", "Note"],
               [list(b) for b in bottleneck],
               col_widths=[Inches(0.9), Inches(3.2), Inches(1.2), Inches(1.5)])


# ─────────────────────────────────────────────
# REPORT 3: Team Workload Report
# ─────────────────────────────────────────────

def build_report3(doc):
    add_report_header(doc, "3. Team Workload Report — May 13, 2026")
    add_body(doc, "Active work distribution across the RIC team. Counts reflect open, non-backlog items only.")

    add_heading(doc, "Workload by Person")
    workload = [
        ("Kyle Hogan",      "5", "5", "4",  "14", "1", "0"),
        ("Thomas Mason",    "3", "7", "6",  "16", "0", "0"),
        ("Aaron West",      "1", "6", "9",  "16", "1", "1"),
        ("Matt Wartell",    "1", "2", "4",  "7",  "2", "2"),
        ("Fareez Ahmed",    "0", "0", "5",  "5",  "0", "0"),
        ("Chris Pusinelli", "0", "1", "2",  "3",  "0", "0"),
        ("Lada Semicheva",  "1", "0", "1",  "2",  "0", "0"),
        ("John Fukuda",     "0", "1", "2",  "3",  "0", "0"),
        ("Ben Higgins",     "1", "0", "0",  "1",  "0", "0"),
        ("Carl Tegeder",    "0", "0", "2",  "2",  "0", "0"),
        ("David Greene",    "0", "0", "2",  "2",  "0", "0"),
        ("Rajesh Pandian",  "0", "0", "1",  "1",  "0", "0"),
        ("Rachel Harrison", "0", "1", "0",  "1",  "0", "0"),
    ]
    make_table(doc,
               ["Person", "Initiatives", "Epics", "Stories/Tasks/Spikes", "Total Active", "Blocked", "Waiting"],
               [list(w) for w in workload],
               col_widths=[Inches(1.3), Inches(0.85), Inches(0.7), Inches(1.35), Inches(0.9), Inches(0.7), Inches(0.7)])

    add_heading(doc, "Observations")
    obs = [
        "Thomas Mason and Aaron West carry the heaviest loads at 16 items each across all work levels.",
        "Matt Wartell has 2 blocked epics — both security use cases (RIC-250, RIC-251) stalled pending client access.",
        "7 unassigned initiatives need owners; proactive assignment recommended.",
        "Kyle Hogan owns 14 items across all hierarchy levels, including 1 blocked initiative (RIC-261, T-MSIS).",
        "Fareez Ahmed and Chris Pusinelli have lighter loads — potential capacity for additional assignments.",
    ]
    for o in obs:
        add_bullet(doc, o)


# ─────────────────────────────────────────────
# REPORT 4: Risk & Blockers Report
# ─────────────────────────────────────────────

def build_report4(doc):
    add_report_header(doc, "4. Risk & Blockers Report — May 13, 2026")
    add_body(doc, "All items that are blocked, waiting, or represent risk to delivery. Built for leadership sync and status emails.")

    add_heading(doc, "Blocked Items (4)")
    blocked = [
        ("RIC-261", "Create Apex Package - T-MSIS",           "Initiative", "Kyle Hogan",  "Critical", "~29 days"),
        ("RIC-251", "Security Use Case 2 - SIA Automation",   "Epic",       "Matt Wartell", "High",    "~35 days"),
        ("RIC-250", "Security Use Case 1 - Scans Review",     "Epic",       "Matt Wartell", "Low",     "~35 days"),
        ("RIC-185", "Collect FOE Success Story from OPERAS",  "Story",      "Aaron West",  "Low",      "~51 days"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Owner", "Priority", "Days Blocked"],
               [list(b) for b in blocked],
               col_widths=[Inches(0.8), Inches(2.5), Inches(0.8), Inches(1.0), Inches(0.75), Inches(0.9)],
               blocked_col=0,
               blocked_values={"RIC-261", "RIC-251", "RIC-250", "RIC-185"})

    add_heading(doc, "Waiting Items (3)")
    waiting = [
        ("RIC-274", "Implement Context Extraction Module",      "Story", "Matt Wartell", "Apr 20"),
        ("RIC-188", "Collect FOE Success Story from USGS",      "Story", "Aaron West",   "Mar 23"),
        ("RIC-407", "Get overview of SIA process from client",  "Story", "Matt Wartell", "May 11"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Owner", "Waiting Since"],
               [list(w) for w in waiting],
               col_widths=[Inches(0.8), Inches(2.8), Inches(0.8), Inches(1.0), Inches(1.0)])

    add_heading(doc, "Blocker/Critical Priority — Not Yet In Progress")
    critical = [
        ("RIC-309", "SHINE MVP - Milestone 2",                     "Epic",  "Kyle Hogan",       "To Do"),
        ("RIC-310", "SHINE MVP - Milestone 3",                     "Epic",  "Unassigned",        "To Do"),
        ("RIC-392", "Defect Management PMR Report Automation",     "Epic",  "Rachel Harrison",   "To Do"),
        ("RIC-420", "Coordinate and support access for test users","Story", "Kyle Hogan",        "To Do"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Owner", "Status"],
               [list(c) for c in critical],
               col_widths=[Inches(0.8), Inches(3.0), Inches(0.8), Inches(1.1), Inches(0.8)])

    add_heading(doc, "Unassigned High/Critical Items")
    unassigned_hi = [
        ("RIC-432", "Support Apex Agent Roll Out For NHLBI",  "Epic", "High",   "Backlog"),
        ("RIC-439", "AWS Partnership Tier Advancement",        "Epic", "High",   "Backlog"),
        ("RIC-438", "Databricks Partnership Tier Advancement", "Epic", "High",   "Backlog"),
        ("RIC-310", "SHINE MVP - Milestone 3",                 "Epic", "High",   "To Do"),
        ("RIC-309", "SHINE MVP - Milestone 2",                 "Epic", "High",   "To Do (Kyle assigned — see note)"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Type", "Priority", "Status"],
               [list(u) for u in unassigned_hi],
               col_widths=[Inches(0.8), Inches(2.8), Inches(0.8), Inches(0.75), Inches(1.75)])

    add_heading(doc, "Stale High-Priority Items (No Update > 14 Days)")
    stale = [
        ("RIC-142", "Agent Orchestration Platform Work",       "Aaron West",   "Apr 7",  "36 days"),
        ("RIC-225", "Katalyst",                                "Thomas Mason", "Apr 17", "26 days"),
        ("RIC-292", "Prima Delivery",                          "Thomas Mason", "Apr 27", "16 days"),
        ("RIC-198", "Get ELT Sign-Off for FOE",                "Aaron West",   "Apr 27", "16 days"),
        ("RIC-297", "White Paper and Technical Documentation",  "John Fukuda",  "Apr 27", "16 days"),
    ]
    make_table(doc,
               ["Ticket", "Summary", "Owner", "Last Updated", "Days Since Update"],
               [list(s) for s in stale],
               col_widths=[Inches(0.8), Inches(2.6), Inches(1.0), Inches(1.0), Inches(1.3)])


# ─────────────────────────────────────────────
# REPORT 5: Monthly Executive Summary
# ─────────────────────────────────────────────

def build_report5(doc):
    add_report_header(doc, "5. Monthly Executive Summary — May 2026")

    add_heading(doc, "Executive Summary")
    exec_text = (
        "The Rapid Innovation Center delivered 5 completed initiatives in the past 30 days, all tied directly "
        "to active business development pursuits and client-facing innovation. The team closed a Databricks POC "
        "series for QDAS, delivered three recompete and proposal assets for CMS MESH, SBA OIG, and the State of "
        "New Mexico, and continued advancing 14 active initiatives. The RIC enters this period with strong "
        "momentum: 8 new initiatives launched, the SHINE platform advancing through 6 MVP milestones, Apex agents "
        "in deployment, and a new partnership management initiative underway."
    )
    add_body(doc, exec_text, bold=True)

    add_heading(doc, "Completed Initiatives — Last 30 Days (5 Items)")

    completions = [
        (
            "1. New Mexico RFP Demo",
            "May 12, 2026  |  Owner: Aaron West",
            "Live modernization demo for NM water rights document management. "
            "Positioned eSimplicity for potential new state government work."
        ),
        (
            "2. DBX Apps POC - Use Case 4",
            "May 8, 2026  |  Owner: Matt Wartell",
            "Proved Databricks automatic app deployment capabilities for non-technical QDAS users. "
            "Directly addresses adoption gap and reduces data scientist bottlenecks."
        ),
        (
            "3. DBX Apps POC - Use Case 3",
            "May 1, 2026  |  Owner: Lada Semicheva",
            "Validated CCSQ Portal running in Databricks Apps microservices architecture. "
            "Informs QDAS roadmap and Databricks partnership investment."
        ),
        (
            "4. SBA OIG - AI Innovation for Recompete",
            "Apr 24, 2026  |  Owner: Kyle Hogan",
            "AI prototype agent for fraud investigators with auditable SQL generation. "
            "High-stakes recompete positioning with a receptive, AI-bullish IG."
        ),
        (
            "5. CMS MESH - Salesforce/M365 Integration",
            "Apr 17, 2026  |  Owner: Kyle Hogan",
            "Closed proposal-critical capability gap with Outlook/Teams integration. "
            "Table-stakes requirement delivered on time."
        ),
    ]

    for name, meta, detail in completions:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.font.name  = BODY_FONT
        r.font.size  = Pt(11)
        r.font.bold  = True
        r.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(0)

        p2 = doc.add_paragraph()
        r2 = p2.add_run(meta)
        r2.font.name   = BODY_FONT
        r2.font.size   = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = RGBColor(100, 100, 100)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)

        p3 = doc.add_paragraph()
        r3 = p3.add_run(detail)
        r3.font.name = BODY_FONT
        r3.font.size = Pt(11)
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after  = Pt(6)

    add_heading(doc, "Portfolio Metrics")
    metrics = [
        ("Total Initiatives (All Time)",      "35"),
        ("Completed",                          "11"),
        ("Active (In Progress / Always On)",   "10"),
        ("Discovery / Ready for Work",         "5"),
        ("To Do / Backlog",                    "4"),
        ("Blocked",                            "1"),
        ("Completed in Last 30 Days",          "5"),
        ("New Initiatives Started (Last 30 Days)", "8"),
        ("All-Time Completion Rate",           "31%"),
    ]
    make_table(doc,
               ["Metric", "Count"],
               [list(m) for m in metrics],
               col_widths=[Inches(3.5), Inches(1.5)])

    add_heading(doc, "Active Highlights")
    highlights = [
        ("SHINE Platform:", "6 milestones planned, M1 in progress, M2–M6 queued."),
        ("Apex Agents:", "Deployment support underway for T-MSIS and CIO Growth."),
        ("Agent Ranker:", "Critical development for NHLBI in progress (RIC-402)."),
        ("Databricks Partnership Day:", "July 8, 2026 event in active planning."),
        ("Growth Opportunities Dashboard:", "In progress — supporting BD pipeline visibility."),
    ]
    for prefix, detail in highlights:
        add_bullet(doc, detail, bold_prefix=prefix + "  ")


# ─────────────────────────────────────────────
# REPORT 6: Domain / Label Distribution
# ─────────────────────────────────────────────

def build_report6(doc):
    add_report_header(doc, "6. Domain & Label Distribution Report — May 13, 2026")
    add_body(doc, "Where is the RIC's energy going? Distribution of active work by business domain and category label.")

    add_heading(doc, "By Domain Label — Initiatives Board (Active, 21 Items)")
    add_body(doc, "Note: Totals exceed 21 because items can carry multiple domain labels.", italic=True)
    domain = [
        ("Health",    "9",  "43%", "T-MSIS Security, NHLBI Copilot, CQP, QDAS, PMR Testing"),
        ("Corporate", "5",  "24%", "SHINE, Flow Optimized Engineering, Partnership Management"),
        ("BD/Growth", "4",  "19%", "Growth Dashboard, IDR Research, Apex CIO Growth"),
        ("DNS",       "2",  "10%", "NTIA Reports, NIWDC AI Agents"),
        ("RIC",       "3",  "14%", "RIC Exploration, FOE, SHINE"),
    ]
    make_table(doc,
               ["Domain", "Count", "% of Board", "Key Initiatives"],
               [list(d) for d in domain],
               col_widths=[Inches(1.0), Inches(0.7), Inches(0.9), Inches(4.2)])

    add_heading(doc, "By Category Label — All Active Work")
    category = [
        ("Program",    "8",  "2", "0",  "10"),
        ("RIC",        "3",  "2", "0",  "5"),
        ("SHINE",      "1",  "4", "4",  "9"),
        ("CIO_Growth", "0",  "0", "5",  "5"),
        ("BD/Growth",  "4",  "3", "1",  "8"),
        ("Corporate",  "5",  "4", "0",  "9"),
        ("FOE",        "0",  "0", "1",  "1"),
    ]
    make_table(doc,
               ["Category", "Initiatives", "Epics", "Stories", "Total"],
               [list(c) for c in category],
               col_widths=[Inches(1.2), Inches(0.9), Inches(0.7), Inches(0.7), Inches(0.7)])

    add_heading(doc, "Observations")
    obs = [
        "Health dominates at 43% of initiatives — appropriate given T-MSIS, NHLBI, QDAS, and CMS programs.",
        "SHINE is cross-cutting: appears in Initiatives, Epics, and Stories — reflecting active build investment.",
        "BD/Growth has strong initiative presence (4) but limited story-level execution — BD work may be tracked at higher levels only.",
        "DNS is underrepresented at 10% — only 2 active initiatives (NTIA, NIWDC). Consider whether this reflects pipeline reality or a tracking gap.",
    ]
    for o in obs:
        add_bullet(doc, o)


# ─────────────────────────────────────────────
# REPORT 7: Initiative Progress Report
# ─────────────────────────────────────────────

def add_initiative_section(doc, emoji_status, title, meta, epics_data, progress_note):
    """Build one initiative section with epic table."""
    p = doc.add_paragraph()
    r = p.add_run(f"{emoji_status} {title}")
    r.font.name  = BODY_FONT
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(meta)
    r2.font.name   = BODY_FONT
    r2.font.size   = Pt(10)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(90, 90, 90)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)

    if epics_data:
        make_table(doc,
                   ["Epic", "Status", "Owner"],
                   [list(e) for e in epics_data],
                   col_widths=[Inches(3.2), Inches(1.3), Inches(1.5)],
                   blocked_col=1,
                   blocked_values={"BLOCKED"})

    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Progress: {progress_note}")
    r3.font.name   = BODY_FONT
    r3.font.size   = Pt(10)
    r3.font.italic = True
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(6)


def build_report7(doc):
    add_report_header(doc, "7. Initiative Progress Report — May 13, 2026")
    add_body(doc, "Status of each active initiative with its associated epics. Shows delivery depth.")

    add_initiative_section(doc,
        "🟡", "SHINE (RIC-340) — Always On | Owner: Kyle Hogan",
        "Cross-cutting platform initiative for the RIC. 6 MVP milestones planned.",
        [
            ("RIC-308  SHINE MVP - Milestone 1", "In Progress", "Kyle Hogan"),
            ("RIC-309  SHINE MVP - Milestone 2", "To Do",       "Kyle Hogan"),
            ("RIC-310  SHINE MVP - Milestone 3", "To Do",       "Unassigned"),
            ("RIC-455  SHINE MVP - Milestone 4", "Backlog",     "Unassigned"),
            ("RIC-457  SHINE MVP - Milestone 5", "Backlog",     "Unassigned"),
            ("RIC-458  SHINE MVP - Milestone 6", "Backlog",     "Unassigned"),
        ],
        "1 of 6 milestones active. M1 in flight, M2–M6 queued."
    )

    add_initiative_section(doc,
        "🟡", "eSim Partnership Management (RIC-436) — Always On | Owner: Kyle Hogan",
        "New initiative (created May 12). Covers vendor partnership tier advancements and events.",
        [
            ("RIC-437  Databricks Partnership Day — July 8", "In Progress", "Kyle Hogan"),
            ("RIC-438  Databricks Tier Advancement (Bronze→Gold)", "Backlog",  "Unassigned"),
            ("RIC-439  AWS Partnership Tier Advancement",          "Backlog",  "Unassigned"),
            ("RIC-440  Salesforce Partnership Tier Advancement",   "Backlog",  "Unassigned"),
            ("RIC-441  Snowflake Partnership Tier Advancement",    "Backlog",  "Unassigned"),
        ],
        "1 of 5 epics active. July 8 event driving immediate activity."
    )

    add_initiative_section(doc,
        "🟡", "Flow Optimized Engineering (RIC-143) — Always On | Owner: Aaron West",
        "Firm-wide FOE adoption program. ELT sign-off gates broader rollout.",
        [
            ("RIC-178  FOE Marketing Across eSimplicity Programs", "In Progress", "Aaron West"),
            ("RIC-180  Implement FOE Public Website",              "In Progress", "Aaron West"),
            ("RIC-198  Get ELT Sign-Off for FOE Implementation",   "In Progress", "Aaron West"),
            ("RIC-175  Publish FOE Book",                          "To Do",       "Unassigned"),
            ("RIC-176  Implement FOE on RIC",                      "To Do",       "Unassigned"),
            ("RIC-177  Pilot FOE at OPERAS",                       "To Do",       "Unassigned"),
            ("RIC-179  Implement FOE on USGS",                     "To Do",       "Unassigned"),
            ("RIC-181  Implement FOE Field Guide",                 "To Do",       "Unassigned"),
            ("RIC-182  Build FOE Scanner",                         "To Do",       "Unassigned"),
        ],
        "3 of 9 epics active. ELT sign-off (RIC-198) is the gate — no update in 16 days."
    )

    add_initiative_section(doc,
        "🟡", "RIC Exploration & Experimentation (RIC-15) — Always On | Owner: Thomas Mason",
        "Ongoing exploratory initiative. Active epics include Tommy's Exploration (RIC-362) and Agent Orchestration Platform Work (RIC-142).",
        [],
        "Continuous. No fixed scope — feeds pipeline and capability development."
    )

    add_initiative_section(doc,
        "🔴", "Create Apex Package - T-MSIS (RIC-261) — BLOCKED | Owner: Kyle Hogan",
        "BLOCKED — Resolution needed to unblock delivery.",
        [
            ("RIC-347  Apex MVP Delivery",                 "In Progress", "Kyle Hogan"),
            ("RIC-418  Support Apex Proposal Deployment",  "In Progress", "Kyle Hogan"),
        ],
        "2 epics active but initiative itself is BLOCKED (~29 days). Immediate attention required."
    )

    add_initiative_section(doc,
        "🟡", "Growth Opportunities Report/Dashboard (RIC-206) — In Progress | Owner: Lada Semicheva",
        "BD pipeline visibility tooling. 1.0 in flight, 2.0 queued.",
        [
            ("RIC-283  Growth Opportunities Reporting 1.0",       "In Progress", "Thomas Mason"),
            ("RIC-284  Growth Portal 2.0 - Agentic AI Ranking",   "To Do",       "Unassigned"),
        ],
        "1.0 in flight, 2.0 queued. No child stories visible on team board."
    )

    add_initiative_section(doc,
        "🟡", "T-MSIS - Security Automation w/ AI (RIC-194) — In Progress | Owner: Matt Wartell",
        "Both epics currently blocked. Active stories waiting on client access.",
        [
            ("RIC-251  Security Use Case 2 - SIA Automation", "BLOCKED", "Matt Wartell"),
            ("RIC-250  Security Use Case 1 - Scans Review",   "BLOCKED", "Matt Wartell"),
        ],
        "⚠️ Both epics are BLOCKED. Stories RIC-273 and RIC-407 are waiting on client access."
    )

    add_initiative_section(doc,
        "🟡", "Create Apex Package - CIO Growth (RIC-417) — In Progress | Owner: Kyle Hogan",
        "New this week (created May 12). No child epics yet established.",
        [],
        "Early stage. Epic RIC-437 (Databricks Partnership Day) is related. Epics to be created."
    )

    add_initiative_section(doc,
        "🟡", "NHLBI - Build Agent for Evaluating Apps for Copilot (RIC-204) — In Progress | Owner: Ben Higgins",
        "Research and ranker development both active. RIC-403 (define ranking criteria) is Blocker priority.",
        [
            ("RIC-224  NHLBI Use Case Research",   "In Progress", "Chris Pusinelli"),
            ("RIC-402  Agent Ranker Development",  "In Progress", "Thomas Mason"),
        ],
        "Research and ranker development both active. RIC-403 is Blocker priority — needs resolution."
    )

    add_initiative_section(doc,
        "🟡", "Assist Health BU with PMR Report Automation Testing (RIC-389) — In Progress | Owner: Kyle Hogan",
        "Testing in flight; defect management epic not yet started.",
        [
            ("RIC-391  HBU PMR PowerApps Testing",       "In Progress", "Kyle Hogan"),
            ("RIC-392  Defect Management and Resolution", "To Do",       "Rachel Harrison"),
        ],
        "Testing in flight. Defect management (RIC-392) not yet started."
    )

    add_initiative_section(doc,
        "🟡", "IDR - Research POC/White Paper (RIC-268) — Reviewing | Owner: Thomas Mason",
        "Blocker priority. In review — needs resolution.",
        [
            ("RIC-297  White Paper and Technical Documentation", "In Progress", "John Fukuda"),
        ],
        "In review. Blocker priority — no update on RIC-297 in 16 days."
    )

    # Discovery stage
    add_heading(doc, "Discovery Stage Initiatives")
    discovery = [
        ("RIC-431", "NHLBI - Apex Agent Assistance",  "Unassigned", "1 day"),
        ("RIC-365", "NIWDC - AI Agents",               "Unassigned", "9 days"),
        ("RIC-316", "T-MSIS - Test Automation",        "Unassigned", "21 days"),
    ]
    make_table(doc,
               ["Initiative", "Summary", "Owner", "Age"],
               [list(d) for d in discovery],
               col_widths=[Inches(0.9), Inches(3.0), Inches(1.1), Inches(0.8)])

    # Ready for Work
    add_heading(doc, "Ready for Work — No Active Epics")
    ready = [
        ("RIC-9",  "CQP - Training Material Automation", "Unassigned", "96 days"),
        ("RIC-10", "QDAS - AI Acceleration",              "Unassigned", "96 days"),
    ]
    make_table(doc,
               ["Initiative", "Summary", "Owner", "Age"],
               [list(r) for r in ready],
               col_widths=[Inches(0.9), Inches(3.0), Inches(1.1), Inches(0.8)])

    add_body(doc,
        "⚠️  Both RIC-9 and RIC-10 have been in Ready for Work for 96 days with no assigned owner. "
        "Leadership review recommended.",
        bold=True, color=RED_COLOR)


# ─────────────────────────────────────────────
# IMPORT STATEMENT FIX (break type)
# ─────────────────────────────────────────────

# Fix add_page_break to use correct import
from docx.enum.text import WD_BREAK

def add_page_break_v2(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup — US Letter, normal margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # Default style
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    # Cover page
    build_cover_page(doc)
    add_page_break_v2(doc)

    # Report 1
    build_report1(doc)
    add_page_break_v2(doc)

    # Report 2
    build_report2(doc)
    add_page_break_v2(doc)

    # Report 3
    build_report3(doc)
    add_page_break_v2(doc)

    # Report 4
    build_report4(doc)
    add_page_break_v2(doc)

    # Report 5
    build_report5(doc)
    add_page_break_v2(doc)

    # Report 6
    build_report6(doc)
    add_page_break_v2(doc)

    # Report 7
    build_report7(doc)

    # Footer
    add_footer(doc)

    # Save
    out_path = "/Users/kyle.hogan/Projects/kyles-arena/RIC-PM-Reports-May2026.docx"
    doc.save(out_path)
    print(f"Document saved: {out_path}")

    import os
    size = os.path.getsize(out_path)
    print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
